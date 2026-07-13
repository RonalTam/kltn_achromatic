from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Bao_cao_tien_do_du_an_Achromatic.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(35, 35, 35)
MUTED = RGBColor(90, 90, 90)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D7DBE2"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_text(cell, text, bold=False, color=INK, size=10.5, align=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    p.text = ""
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_table(doc, headers, rows, widths_dxa, header_fill=LIGHT_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    mark_header_row(table.rows[0])
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], header_fill)
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i in (0, len(row) - 1) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], str(value), align=align)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    for run in p.runs:
        set_run_font(run, size=11, color=INK)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.line_spacing = 1.1
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT_FILL)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "4")
        node.set(qn("w:color"), "E1E5EA")
        borders.append(node)
    p_pr.append(borders)
    r1 = p.add_run(f"{label}: ")
    set_run_font(r1, size=11, color=DARK_BLUE, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=11, color=INK)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_metadata(doc):
    rows = [
        ("Đơn vị nhận", "Nhà trường / Khoa chuyên môn"),
        ("Tên dự án", "Achromatic Fashion E-Commerce Platform"),
        ("Nhóm / Sinh viên", "........................................................"),
        ("Giảng viên hướng dẫn", "........................................................"),
        ("Ngày báo cáo", "07/07/2026"),
        ("Trạng thái tổng quát", "MVP đã hoàn thành, đang chuyển sang giai đoạn hoàn thiện UX, kiểm thử và triển khai"),
    ]
    for label, value in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}: ")
        set_run_font(r1, size=11, color=INK, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=11, color=INK)


def build():
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = ""
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = header_p.add_run("Báo cáo tiến độ dự án | Achromatic Fashion E-Commerce")
    set_run_font(hr, size=9, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run("Tài liệu báo cáo nội bộ phục vụ đánh giá tiến độ")
    set_run_font(fr, size=9, color=MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(4)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("BÁO CÁO TIẾN ĐỘ DỰ ÁN")
    set_run_font(tr, size=22, color=RGBColor(0, 0, 0), bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    sr = subtitle.add_run("Xây dựng website thương mại điện tử thời trang Achromatic")
    set_run_font(sr, size=13, color=MUTED, italic=True)

    add_metadata(doc)
    add_callout(
        doc,
        "Tóm tắt",
        "Dự án đã hoàn thành phần lõi của một hệ thống thương mại điện tử full-stack: giao diện mua sắm, danh mục sản phẩm, giỏ hàng, thanh toán, đăng nhập/đăng ký, hệ thống API, cơ sở dữ liệu PostgreSQL và tài liệu kỹ thuật. Kết quả build ngày 07/07/2026 cho cả frontend và backend đều thành công.",
    )

    add_heading(doc, "1. Mục tiêu và phạm vi dự án", 1)
    add_body(
        doc,
        "Dự án hướng tới xây dựng một website thương mại điện tử thời trang theo phong cách tối giản, cho phép người dùng xem sản phẩm, lọc danh mục, xem chi tiết sản phẩm, quản lý giỏ hàng, đăng nhập tài khoản và thực hiện quy trình checkout cơ bản. Phía quản trị và các chức năng nâng cao được chuẩn bị nền tảng ở backend để tiếp tục mở rộng.",
    )

    add_heading(doc, "2. Công nghệ sử dụng", 1)
    tech_rows = [
        ("Frontend", "Next.js 16.2.9, React 19.2.4, TypeScript 5, TailwindCSS 4, shadcn/ui, Zustand, Axios, React Query"),
        ("Backend", "NestJS 11, TypeScript, Prisma 7.8, PostgreSQL, Passport JWT/Local, Swagger/OpenAPI"),
        ("Cơ sở dữ liệu", "Schema Prisma với các nhóm bảng cho người dùng, sản phẩm, biến thể, giỏ hàng, đơn hàng, thanh toán, đánh giá, wishlist, banner, blog, thông báo và audit log"),
        ("Công cụ hỗ trợ", "npm, Git, Prisma migration/seed, Swagger UI, script kiểm thử API"),
    ]
    add_table(doc, ["Mảng", "Công nghệ / Ghi chú"], tech_rows, [1800, 7560])

    add_heading(doc, "3. Tiến độ thực hiện theo hạng mục", 1)
    progress_rows = [
        ("Phân tích & thiết kế", "Đã xác định phạm vi, luồng người dùng, cấu trúc frontend/backend và mô hình dữ liệu chính.", "90%", "Hoàn thành cơ bản"),
        ("Cơ sở dữ liệu", "Đã xây dựng schema Prisma, migration khởi tạo và seed dữ liệu mẫu gồm sản phẩm, thương hiệu, danh mục, banner, tài khoản test.", "85%", "Cần bổ sung dữ liệu thật khi triển khai"),
        ("Backend API", "Đã có các module chính: auth, users, products, categories, brands, cart, orders, payments, shipping, reviews, wishlists, collections, banners, blogs, notifications, analytics, admin, inventory, coupons.", "80%", "Cần kiểm thử sâu và hoàn thiện một số luồng nghiệp vụ"),
        ("Frontend UI", "Đã có trang chủ, collections, chi tiết sản phẩm, giỏ hàng, checkout, đăng nhập/đăng ký, tài khoản, wishlist, chính sách, liên hệ, FAQ, admin skeleton.", "80%", "Cần polish UX và trạng thái loading/error"),
        ("Tích hợp API", "Frontend đã tích hợp các luồng sản phẩm, danh mục, giỏ hàng và xác thực ở mức MVP.", "75%", "Wishlist/review/order tracking cần đồng bộ đầy đủ hơn"),
        ("Kiểm thử", "Đã build thành công frontend và backend ngày 07/07/2026; tài liệu cũ ghi nhận các API chính trả về 200 OK.", "60%", "Cần kiểm thử trình duyệt, e2e và test case chi tiết"),
        ("Tài liệu", "Có README, QUICK_START, PROGRESS, TASKS, DATABASE_STATUS, FINAL_STATUS và script hỗ trợ test API.", "75%", "Cần chuẩn hóa lại encoding tiếng Việt ở một số markdown"),
        ("Triển khai", "Chưa triển khai production; đã có cấu trúc phù hợp để chuẩn bị deploy.", "25%", "Cần cấu hình môi trường, domain, CI/CD và monitoring"),
    ]
    add_table(doc, ["Hạng mục", "Kết quả hiện tại", "Ước lượng", "Ghi chú"], progress_rows, [1550, 4350, 1150, 2310])

    add_heading(doc, "4. Các kết quả đã hoàn thành", 1)
    add_heading(doc, "4.1. Frontend", 2)
    for item in [
        "Xây dựng giao diện chính bằng Next.js App Router, TypeScript và TailwindCSS.",
        "Hoàn thành các màn hình cốt lõi: trang chủ, danh sách sản phẩm, chi tiết sản phẩm, giỏ hàng, checkout, đăng nhập, đăng ký, tài khoản, đơn hàng, cài đặt, wishlist.",
        "Bổ sung các trang nội dung: giới thiệu, liên hệ, FAQ, hướng dẫn kích cỡ, chính sách vận chuyển, đổi trả, bảo mật và điều khoản.",
        "Xây dựng component dùng lại: Header, Footer, ProductCard, ProductDetail, FilterPanel, SortDropdown, Pagination, CartDrawer, HeroSlider, Newsletter và các section trang chủ.",
        "Thiết lập store cho xác thực và giỏ hàng bằng Zustand, có lưu trạng thái phía client.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4.2. Backend", 2)
    for item in [
        "Thiết lập ứng dụng NestJS với cấu trúc module rõ ràng theo nghiệp vụ.",
        "Hoàn thiện các nhóm API chính cho xác thực, người dùng, sản phẩm, danh mục, thương hiệu, giỏ hàng, đơn hàng, thanh toán, vận chuyển, đánh giá, wishlist, bộ sưu tập, banner, blog, thông báo, phân tích và quản trị.",
        "Tích hợp Prisma với PostgreSQL, có schema cho hơn 40 bảng nghiệp vụ.",
        "Hỗ trợ JWT authentication, guard phân quyền, validation DTO và Swagger/OpenAPI để tra cứu API.",
        "Có seed data và tài khoản test phục vụ demo.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4.3. Dữ liệu mẫu và nội dung demo", 2)
    data_rows = [
        ("Sản phẩm", "25 sản phẩm mẫu thuộc nhiều nhóm thời trang"),
        ("Danh mục", "5 danh mục chính: Tops, Bottoms, Outerwear, Accessories, Shoes"),
        ("Thương hiệu", "10 thương hiệu mẫu"),
        ("Biến thể", "Khoảng 450+ biến thể màu/size theo tài liệu seed"),
        ("Banner & collection", "3 hero banners và 3 collections phục vụ trang chủ"),
        ("Tài khoản test", "Admin, Manager và Customer để phục vụ kiểm thử/demo"),
    ]
    add_table(doc, ["Nhóm dữ liệu", "Trạng thái"], data_rows, [2200, 7160])

    add_heading(doc, "5. Kết quả kiểm tra hiện tại", 1)
    verify_rows = [
        ("Frontend production build", "Đạt", "Đã chạy `npm.cmd run build` trong thư mục frontend; Next.js compile, TypeScript và static generation thành công."),
        ("Backend build", "Đạt", "Đã chạy `npm.cmd run build` trong thư mục backend; NestJS build thành công."),
        ("API cơ bản", "Đã ghi nhận trong tài liệu", "Các endpoint sản phẩm, danh mục, thương hiệu và banner được tài liệu nội bộ ghi nhận trả về 200 OK vào 17/06/2026."),
        ("Kiểm thử trình duyệt", "Cần bổ sung", "Chưa thực hiện lại manual browser test trong lần lập báo cáo này; cần kiểm tra các luồng mua hàng end-to-end."),
    ]
    add_table(doc, ["Nội dung kiểm tra", "Kết quả", "Ghi chú"], verify_rows, [2600, 1450, 5310])

    add_heading(doc, "6. Hạn chế và vấn đề còn tồn tại", 1)
    for item in [
        "Chưa có toast notification, skeleton loading và error boundary đầy đủ.",
        "Chức năng tìm kiếm toàn cục, gửi đánh giá, chi tiết theo dõi đơn hàng và đồng bộ wishlist với backend chưa hoàn thiện.",
        "Giao diện quản trị mới ở mức nền tảng/skeleton, chưa đủ chức năng quản lý sản phẩm, đơn hàng, người dùng và thống kê.",
        "Chưa tối ưu sâu SEO, accessibility, ảnh bằng next/image, bundle size và triển khai production.",
        "Một số file tài liệu markdown đang bị lỗi encoding tiếng Việt, cần chuẩn hóa để nộp kèm hồ sơ dự án.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. Kế hoạch thực hiện tiếp theo", 1)
    plan_rows = [
        ("Tuần 1", "Hoàn thiện UX", "Toast notification, loading states, thông báo lỗi, kiểm thử responsive/mobile."),
        ("Tuần 2", "Tính năng người dùng", "Tìm kiếm sản phẩm, review submission, wishlist sync, trang chi tiết đơn hàng và tracking."),
        ("Tuần 3", "Quản trị", "Dashboard admin, quản lý sản phẩm, đơn hàng, người dùng và báo cáo thống kê."),
        ("Tuần 4", "Production", "SEO, performance, accessibility, biến môi trường production, deploy frontend/backend/database và tài liệu hướng dẫn sử dụng."),
    ]
    add_table(doc, ["Giai đoạn", "Trọng tâm", "Công việc chính"], plan_rows, [1300, 2100, 5960])

    add_heading(doc, "8. Đánh giá chung và kiến nghị", 1)
    add_body(
        doc,
        "Nhìn chung, dự án đã đạt mức MVP: các luồng mua sắm và nền tảng kỹ thuật chính đã được xây dựng, có dữ liệu mẫu và có thể tiếp tục demo sau khi chạy môi trường phát triển. Nhóm đề xuất tiếp tục được triển khai giai đoạn hoàn thiện để nâng chất lượng trải nghiệm người dùng, kiểm thử đầy đủ hơn và chuẩn bị triển khai thực tế.",
    )
    add_body(
        doc,
        "Kiến nghị nhà trường/giảng viên xem xét ghi nhận tiến độ hiện tại là hoàn thành phần lõi, đồng thời cho phép nhóm tiếp tục phát triển các phần nâng cao như quản trị, tìm kiếm, đánh giá sản phẩm, theo dõi đơn hàng và triển khai production.",
    )

    add_heading(doc, "Phụ lục: nguồn tham chiếu trong dự án", 1)
    for item in [
        "README.md - mô tả tổng quan dự án và hướng dẫn chạy.",
        "PROGRESS.md / FINAL_STATUS.md - ghi nhận tiến độ MVP và các chức năng đã hoàn thành.",
        "DATABASE_STATUS.md - trạng thái dữ liệu mẫu và API đã kiểm tra.",
        "TASKS.md - checklist các phase tiếp theo.",
        "frontend/package.json và backend/package.json - thông tin công nghệ và script build.",
        "backend/prisma/schema.prisma - mô hình dữ liệu hệ thống.",
    ]:
        add_bullet(doc, item)

    doc.core_properties.title = "Báo cáo tiến độ dự án Achromatic Fashion E-Commerce"
    doc.core_properties.subject = "Báo cáo tiến độ gửi nhà trường"
    doc.core_properties.author = "Nhóm thực hiện dự án"
    doc.core_properties.comments = "Generated from local project status on 07/07/2026."
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
